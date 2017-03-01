# _*_ coding: utf_8  _*_
#! python3
# try to write some content to a new Word document.

import logging
import docx

logging.basicConfig( level = logging.DEBUG, format = ' %(asctime)s - %(levelname)s - %(message)s' )
# logging.basicConfig( level = logging.ERROR, format = ' %(asctime)s - %(levelname)s - %(message)s' )


def addruns( filename, runs):
    # Add runs to a paragraph.
    doc = docx.Document(filename)
    for r in runs:
        doc.add_paragraph('  ').add_run(r)
    doc.save( filename)

def writepara( filename, para):
    # Write a paragraph to a new document.
    doc = docx.Document()
    for p in para:
        doc.add_paragraph(p)
    paraobj3 = doc.add_paragraph( 'This is a paraObj3.')
    paraobj4 = doc.add_paragraph( 'This is a paraObj4.')
    paraobj3.add_run('This is being added to the paraObj3')
    paraobj4.add_run('This is being added to the paraObj4')
    doc.save( filename)
    

def main():
    DSTDOC = 'writetestdoc.docx'
    PARA = ['This is a paraObj1 on the first page!', 'This is a paraObj2.']
    writepara( DSTDOC, PARA)
    input()
    RUNS = ['This is being added to the paraObj1', 'This is being added to the paraObj2']
    addruns( DSTDOC, RUNS)
    
##    import getfile_fun
##    DSTDOC = getfile_fun.getfilename('destinationfile')

if __name__ == '__main__':
    logging.critical('--------Begin---------')
    main()
    logging.critical('--------End---------')
    





