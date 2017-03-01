# _*_ coding: utf_8  _*_
#! python3
# read and print the docx successfully. 2017-2-26
# this program can get text in document by several mathods.
# but 若用p=ui算得的功率值为 -_  5W, before 5W a symbol is missing.

import logging
import docx
import getfile_fun

logging.basicConfig( level = logging.DEBUG, format = ' %(asctime)s - %(levelname)s - %(message)s' )
# logging.basicConfig( level = logging.ERROR, format = ' %(asctime)s - %(levelname)s - %(message)s' )
logging.critical('--------Start of program---------')

def getruns(filename, row):
    doc = docx.Document(filename)
    run_ = ''
    for k in range(len(doc.paragraphs[row].runs)):
        print(k, doc.paragraphs[row].runs[k].text, ' ')   #
        run_ += doc.paragraphs[row].runs[k].text
    return run_

def getpara(filename, row):
    doc = docx.Document(filename)
    return doc.paragraphs[row].text

def getparas(filename):
    doc = docx.Document(filename)
    fullparas = []
    for para in doc.paragraphs:
        fullparas.append('  '+para.text)
    return '\n'.join(fullparas)

def main():
##    SRCDOC = 'testDoc-电路ch1.docx'
    SRCDOC = getfile_fun.getfilename('sourcefile')
    paras = getparas(SRCDOC)
    print(paras[0:60])
    print('\n...............')
    for k in range(len(paras)):
        print(paras[k], end='')
    print('\n', k)
    print('\n...............')
    for k in range(11):
        line = getpara(SRCDOC, k)
        print(k, line)
    print('\n...............')
    for k in range(8):
        run = getruns(SRCDOC, k)
        print('\n', k, run, '\n')

if __name__ == '__main__':
    main()




